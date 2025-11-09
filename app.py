import os
import json
import tempfile
import shutil
import subprocess
import base64
import io
import hashlib
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed
from flask import Flask, render_template, request, jsonify, send_file, make_response
from flask_socketio import SocketIO, emit
from datetime import datetime
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import docx.enum.text
from copy import deepcopy
from docxcompose.composer import Composer
from PIL import Image
try:
    from pdf2image import convert_from_path
except:
    convert_from_path = None

try:
    import fitz  # PyMuPDF - szybsza alternatywa do pdf2image
except ImportError:
    fitz = None

# Mutex dla LibreOffice (tylko jedna konwersja naraz)
libreoffice_lock = threading.Lock()

# Mutex dla app4.py/unoserver (tylko jedna konwersja naraz - unoserver single-threaded!)
app4_lock = threading.Lock()

# Cache dla skonwertowanych obrazów (produkty)
conversion_cache = {}

# INTELIGENTNY CACHE dla kompletnych ofert
# Klucz: hash(template_id + form_data + selected_products)
# Wartość: {'pages': [lista obrazów], 'total_pages': N, 'timestamp': T}
offer_cache = {}

# Cache dla poszczególnych stron (fallback)
page_cache = {}

app = Flask(__name__)
app.config['SECRET_KEY'] = 'secret_key_for_socketio_12345'
app.config['COMPRESS_MIMETYPES'] = ['application/json', 'text/html', 'text/css', 'application/javascript']
app.config['COMPRESS_LEVEL'] = 6
app.config['COMPRESS_MIN_SIZE'] = 500

# Włącz kompresję gzip dla wszystkich odpowiedzi
from flask_compress import Compress
Compress(app)

socketio = SocketIO(app, cors_allowed_origins="*")

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATES_DIR = os.path.join(BASE_DIR, 'templates')
PRODUKTY_DIR = os.path.join(BASE_DIR, 'produkty')
SAVED_OFFERS_DIR = os.path.join(BASE_DIR, 'saved_offers')
GENERATED_OFFERS_DIR = os.path.join(BASE_DIR, 'generated_offers')
PREVIEW_CACHE_DIR = os.path.join(BASE_DIR, 'static', 'preview_cache')

# Upewnij się, że foldery istnieją
os.makedirs(SAVED_OFFERS_DIR, exist_ok=True)
os.makedirs(GENERATED_OFFERS_DIR, exist_ok=True)
os.makedirs(PREVIEW_CACHE_DIR, exist_ok=True)


def preload_all_products():
    """Pre-generuje wszystkie produkty przy starcie aplikacji"""
    print("[STARTUP] Pre-generowanie produktów...")

    if not os.path.exists(PRODUKTY_DIR):
        print("[STARTUP] Folder produktów nie istnieje")
        return

    product_files = [f for f in os.listdir(PRODUKTY_DIR) if f.endswith('.docx')]
    total = len(product_files)

    print(f"[STARTUP] Znaleziono {total} produktów do pre-generowania")

    for idx, filename in enumerate(product_files):
        product_path = os.path.join(PRODUKTY_DIR, filename)
        print(f"[STARTUP] Pre-generuję {idx+1}/{total}: {filename}")

        try:
            # Konwertuj z cache - przy pierwszym uruchomieniu wypełni cache
            convert_docx_to_images(product_path, use_cache=True)
            print(f"[STARTUP] ✓ {filename} gotowy")
        except Exception as e:
            print(f"[STARTUP] ✗ Błąd dla {filename}: {e}")

    print(f"[STARTUP] Pre-generowanie zakończone! Cache zawiera {len(conversion_cache)} produktów")


def preload_all_templates():
    """
    NOWA STRATEGIA - NIE RENDERUJ czystych szablonów!

    Dlaczego?
    - Czysty szablon ma {{placeholders}} - i tak trzeba będzie go re-renderować
    - To STRATA CZASU przy starcie
    - Lepiej renderować ON-DEMAND z prawdziwymi danymi

    Co robimy zamiast tego:
    - Pre-renderujemy tylko PRODUKTY (nie zmieniają się)
    - Cache tworzymy PODCZAS pierwszego użycia (z prawdziwymi danymi)
    """
    print("=" * 80)
    print("[STARTUP] 💡 INTELIGENTNY SYSTEM CACHE")
    print("=" * 80)
    print("[STARTUP] ℹ️  Szablony będą renderowane ON-DEMAND (szybsze!)")
    print("[STARTUP] ℹ️  Produkty są pre-renderowane (gotowe do użycia)")
    print("=" * 80)

    # Inicjalizuj pusty cache - wypełni się podczas użytkowania
    # Klucz: hash(template_id + form_data + products)
    # Wartość: [lista obrazów base64]


# Pre-generuj produkty przy starcie (w tle, nie blokuj startu)
def preload_products_async():
    """Uruchom pre-generowanie w osobnym wątku"""
    import threading
    thread = threading.Thread(target=preload_all_products, daemon=True)
    thread.start()

def preload_templates_async():
    """Uruchom pre-generowanie szablonów w osobnym wątku"""
    import threading
    thread = threading.Thread(target=preload_all_templates, daemon=True)
    thread.start()

def get_file_hash(filepath):
    """Oblicza hash pliku dla cache"""
    try:
        with open(filepath, 'rb') as f:
            return hashlib.md5(f.read()).hexdigest()
    except:
        return None


def get_form_data_hash(form_data):
    """Oblicza hash danych formularza dla cache szablonu"""
    try:
        # Sortuj klucze dla stabilnego hasha
        sorted_data = json.dumps(form_data, sort_keys=True)
        return hashlib.md5(sorted_data.encode('utf-8')).hexdigest()
    except:
        return None


def get_offer_cache_key(template_id, form_data, selected_products, product_custom_fields):
    """
    Generuje unikalny klucz cache dla kompletnej oferty

    Klucz uwzględnia:
    - template_id (aidrops/wolftax)
    - form_data (dane formularza - placeholders)
    - selected_products (lista ID produktów)
    - product_custom_fields (custom pola produktów)

    Returns: string (hash MD5)
    """
    try:
        cache_key_data = {
            'template_id': template_id,
            'form_data': form_data,
            'selected_products': sorted(selected_products),  # sort dla stabilności
            'product_custom_fields': product_custom_fields
        }
        cache_key_json = json.dumps(cache_key_data, sort_keys=True)
        cache_key_hash = hashlib.md5(cache_key_json.encode('utf-8')).hexdigest()
        return cache_key_hash
    except Exception as e:
        print(f"[CACHE] ❌ Błąd generowania klucza: {e}")
        return None


def send_progress(message, percent):
    """Wysyła progress przez WebSocket"""
    try:
        socketio.emit('conversion_progress', {
            'message': message,
            'percent': percent
        })
    except:
        pass  # Ignoruj błędy WebSocket


def send_page_ready(page_data):
    """Wysyła gotową stronę przez WebSocket (streaming)"""
    try:
        socketio.emit('page_ready', page_data)
        print(f"[DEBUG] ✓ Wysłano stronę {page_data.get('number')} przez WebSocket")
    except Exception as e:
        print(f"[ERROR] Błąd wysyłania strony: {e}")


def check_unoserver_running():
    """Sprawdza czy unoserver jest uruchomiony"""
    try:
        result = subprocess.run(['pgrep', '-f', 'unoserver'], capture_output=True, text=True)
        return result.returncode == 0
    except:
        return False


def start_unoserver():
    """Uruchamia unoserver w trybie daemon jeśli nie jest uruchomiony"""
    if check_unoserver_running():
        print("[UNOSERVER] ✓ Unoserver już działa")
        return True

    print("[UNOSERVER] Uruchamiam unoserver --daemon...")
    try:
        # Uruchom unoserver w tle
        subprocess.Popen(
            ['unoserver', '--daemon'],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
            start_new_session=True
        )

        # Poczekaj chwilę na start
        import time
        time.sleep(2)

        if check_unoserver_running():
            print("[UNOSERVER] ✓ Unoserver uruchomiony pomyślnie")
            return True
        else:
            print("[UNOSERVER] ❌ Nie udało się uruchomić unoserver")
            return False
    except FileNotFoundError:
        print("[UNOSERVER] ❌ Nie znaleziono unoserver w PATH")
        return False
    except Exception as e:
        print(f"[UNOSERVER] ❌ Błąd uruchamiania: {e}")
        return False


def convert_docx_to_images_unoserver(docx_path, dpi=200, quality=90, uc_host='127.0.0.1', uc_port=2003):
    """
    SUPER FAST DOCX → JPG używając app4.py!
    Wywołuje: python app4.py DOCX -o OUT --dpi DPI --quality Q --uc-host H --uc-port P

    WAŻNE: Unoserver obsługuje tylko 1 konwersję naraz - używamy mutex!
    """
    print(f"[APP4] Czekam na dostęp do unoserver...")

    # MUTEX - tylko jedna konwersja app4.py/unoserver naraz!
    with app4_lock:
        print(f"[APP4] 🚀 START: {os.path.basename(docx_path)}")

        # Sprawdź czy app4.py istnieje
        app4_path = os.path.join(BASE_DIR, 'app4.py')
        if not os.path.exists(app4_path):
            print(f"[APP4] ❌ Brak app4.py w {BASE_DIR}")
            return None

        # Tymczasowy katalog na JPG
        temp_outdir = tempfile.mkdtemp(dir=PREVIEW_CACHE_DIR)

        try:
            # Wywołaj app4.py DOKŁADNIE jak w przykładzie!
            cmd = [
                'python',
                app4_path,
                docx_path,
                '-o', temp_outdir,
                '--dpi', str(dpi),
                '--quality', str(quality),
                '--uc-host', uc_host,
                '--uc-port', str(uc_port)
            ]

            print(f"[APP4] Cmd: {' '.join(cmd)}")

            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=30  # 30s timeout - dużo czasu dla dużych plików
            )

            if result.returncode != 0:
                print(f"[APP4] ❌ Failed:")
                print(f"[APP4] stdout: {result.stdout[:300]}")
                print(f"[APP4] stderr: {result.stderr[:300]}")
                return None

            print(f"[APP4] stdout: {result.stdout.strip()}")

            # Zbierz wygenerowane JPG
            jpg_files = sorted([
                f for f in os.listdir(temp_outdir)
                if f.lower().endswith(('.jpg', '.jpeg'))
            ])

            if not jpg_files:
                print(f"[APP4] ❌ Brak plików JPG w {temp_outdir}")
                return None

            print(f"[APP4] ✓ Znaleziono {len(jpg_files)} plików JPG")

            # Wczytaj JPG i konwertuj na base64
            image_data_list = []
            for jpg_file in jpg_files:
                jpg_path = os.path.join(temp_outdir, jpg_file)
                with open(jpg_path, 'rb') as f:
                    jpg_bytes = f.read()
                    img_base64 = base64.b64encode(jpg_bytes).decode('utf-8')
                    image_data_list.append(f"data:image/jpeg;base64,{img_base64}")

            print(f"[APP4] ✓ SUPER FAST: {len(image_data_list)} stron")
            return image_data_list

        except subprocess.TimeoutExpired:
            print("[APP4] ❌ Timeout (>30s)")
            return None
        except FileNotFoundError as e:
            print(f"[APP4] ❌ File not found: {e}")
            return None
        except Exception as e:
            print(f"[APP4] ❌ Error: {e}")
            import traceback
            traceback.print_exc()
            return None
        finally:
            # Cleanup
            try:
                shutil.rmtree(temp_outdir)
            except:
                pass


def convert_docx_to_images(docx_path, use_cache=True, progress_callback=None):
    """
    Konwertuje plik DOCX na obrazy JPEG (1:1 wygląd)
    Zwraca listę ścieżek do obrazów (base64 encoded)

    Strategia:
    1. Sprawdź cache
    2. Spróbuj unoserver (szybkie!)
    3. Fallback: LibreOffice + pdf2image
    """
    try:
        # Sprawdź cache
        if use_cache:
            file_hash = get_file_hash(docx_path)
            if file_hash and file_hash in conversion_cache:
                print(f"[CACHE] ⚡ Używam cache dla: {docx_path}")
                return conversion_cache[file_hash]

        print(f"[CONVERT] Rozpoczynam konwersję: {docx_path}")

        # ===== SPRÓBUJ APP4.PY (SUPER FAST!) =====
        if check_unoserver_running():
            print(f"[CONVERT] 🚀 Używam app4.py (SUPER FAST)...")
            if progress_callback:
                progress_callback("⚡ SUPER FAST conversion (app4.py)...", 20)

            images = convert_docx_to_images_unoserver(docx_path, dpi=200, quality=90)

            if images:
                # SUKCES! Zapisz do cache
                if use_cache:
                    file_hash = get_file_hash(docx_path)
                    if file_hash:
                        conversion_cache[file_hash] = images
                print(f"[CONVERT] ✓ APP4: {len(images)} stron")
                return images
            else:
                print(f"[CONVERT] ⚠️ APP4 failed, fallback do LibreOffice...")

        # ===== FALLBACK: LIBREOFFICE + pdf2image =====
        print(f"[CONVERT] Fallback: LibreOffice + pdf2image")

        # Najpierw konwertuj DOCX na PDF używając LibreOffice/soffice
        pdf_path = docx_path.replace('.docx', '.pdf')

        # Sprawdź czy LibreOffice jest dostępne
        soffice_paths = [
            '/Applications/LibreOffice.app/Contents/MacOS/soffice',
            '/usr/local/bin/soffice',
            'soffice'
        ]

        soffice_path = None
        for path in soffice_paths:
            if os.path.exists(path) or path == 'soffice':
                soffice_path = path
                print(f"[DEBUG] Znaleziono LibreOffice: {path}")
                break

        if not soffice_path:
            print("[ERROR] LibreOffice nie znalezione!")
            return []

        # MUTEX - tylko jedna konwersja LibreOffice naraz
        with libreoffice_lock:
            try:
                # Konwersja DOCX → PDF
                print(f"[DEBUG] Konwertuję DOCX → PDF: {docx_path}")
                if progress_callback:
                    progress_callback("Konwersja do PDF...", 20)

                result = subprocess.run([
                    soffice_path,
                    '--headless',
                    '--convert-to', 'pdf',
                    '--outdir', os.path.dirname(docx_path),
                    docx_path
                ], check=True, capture_output=True, timeout=60)  # Zwiększono timeout do 60s

                print(f"[DEBUG] LibreOffice stdout: {result.stdout.decode('utf-8', errors='ignore')}")
                print(f"[DEBUG] LibreOffice stderr: {result.stderr.decode('utf-8', errors='ignore')}")

            except subprocess.CalledProcessError as e:
                print(f"[ERROR] LibreOffice błąd: {e}")
                print(f"[ERROR] stdout: {e.stdout.decode('utf-8', errors='ignore')}")
                print(f"[ERROR] stderr: {e.stderr.decode('utf-8', errors='ignore')}")
                return []
            except FileNotFoundError:
                print(f"[ERROR] Nie znaleziono pliku: {soffice_path}")
                return []
            except subprocess.TimeoutExpired:
                print("[ERROR] Timeout podczas konwersji LibreOffice")
                return []

        # Sprawdź czy PDF został utworzony
        if not os.path.exists(pdf_path):
            print(f"[ERROR] PDF nie został utworzony: {pdf_path}")
            return []

        print(f"[DEBUG] PDF utworzony: {pdf_path}")

        # Konwertuj PDF na obrazy PNG
        if not convert_from_path:
            print("[ERROR] pdf2image nie jest zainstalowane!")
            return []

        try:
            print(f"[DEBUG] Konwertuję PDF → PNG")
            if progress_callback:
                progress_callback("Konwersja PDF na obrazy...", 30)

            images = convert_from_path(pdf_path, dpi=100)  # Zmniejszono z 150 na 100 dla szybkości
            print(f"[DEBUG] Wygenerowano {len(images)} stron")

            image_data_list = []
            for i, image in enumerate(images):
                # Progress: Konwersja stron
                if progress_callback and len(images) > 1:
                    page_progress = 30 + int(10 * ((i + 1) / len(images)))
                    progress_callback(f"Przetwarzam stronę {i+1}/{len(images)}...", page_progress)

                # Optymalizuj obraz przed konwersją do base64
                # 1. Konwertuj do RGB (usunąć alpha channel jeśli nie potrzebny)
                if image.mode == 'RGBA':
                    # Tylko konwertuj jeśli nie ma transparentności
                    background = Image.new('RGB', image.size, (255, 255, 255))
                    background.paste(image, mask=image.split()[3])  # 3 to alpha channel
                    image = background

                # 2. Zapisz jako JPEG z kompresją dla mniejszego rozmiaru
                # JPEG jest ~70% mniejszy niż PNG przy zachowaniu dobrej jakości
                buffered = io.BytesIO()
                image.save(buffered, format="JPEG", quality=85, optimize=True)
                img_base64 = base64.b64encode(buffered.getvalue()).decode('utf-8')
                image_data_list.append(f"data:image/jpeg;base64,{img_base64}")
                print(f"[DEBUG] Strona {i+1} skonwertowana (rozmiar: {len(img_base64)} znaków, format: JPEG)")

            # Usuń tymczasowy PDF
            try:
                os.remove(pdf_path)
                print(f"[DEBUG] Usunięto tymczasowy PDF")
            except:
                pass

            # Zapisz w cache
            if use_cache and file_hash:
                conversion_cache[file_hash] = image_data_list
                print(f"[DEBUG] Zapisano w cache: {docx_path}")

            return image_data_list

        except Exception as e:
            print(f"[ERROR] Błąd konwersji PDF → PNG: {e}")
            import traceback
            traceback.print_exc()
            return []

    except Exception as e:
        print(f"[ERROR] Ogólny błąd konwersji DOCX na obrazy: {e}")
        import traceback
        traceback.print_exc()
        return []


def load_template_config(template_name='oferta1'):
    """Wczytuje konfigurację szablonu z pliku JSON"""
    config_path = os.path.join(TEMPLATES_DIR, f'{template_name}.json')
    with open(config_path, 'r', encoding='utf-8') as f:
        return json.load(f)


def extract_placeholders_from_docx(docx_path):
    """
    Wyciąga wszystkie {{placeholders}} z pliku DOCX
    Zwraca set z nazwami placeholders
    """
    import re
    try:
        doc = Document(docx_path)
        placeholders = set()

        # Szukaj w paragrafach
        for para in doc.paragraphs:
            matches = re.findall(r'\{\{([^}]+)\}\}', para.text)
            placeholders.update(matches)

        # Szukaj w tabelach
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    matches = re.findall(r'\{\{([^}]+)\}\}', cell.text)
                    placeholders.update(matches)

        return placeholders
    except Exception as e:
        print(f"[ERROR] Błąd wyciągania placeholders z {docx_path}: {e}")
        return set()


def generate_table_of_contents(selected_products, product_custom_fields, start_page=5):
    """
    Generuje spis treści dla WolfTax w formacie:
    §	Usługa 1 – Nazwa usługi ………….……  05
    """
    toc_lines = []
    current_page = start_page

    for idx, product_id in enumerate(selected_products, 1):
        # Pobierz tytuł z custom fields lub użyj domyślnej nazwy
        product_data = product_custom_fields.get(product_id, {})
        title = product_data.get('title') or product_data.get('nazwa') or product_data.get('nazwa_uslugi') or f'Produkt {product_id}'

        # Oblicz liczbę kropek (dots) dla wyrównania
        # Format: "§\tUsługa 1 – {title} ………….……  {page}"
        base_text = f"Usługa {idx} – {title}"
        dots_count = max(60 - len(base_text), 10)  # Minimum 10 kropek
        dots = '…' * dots_count

        line = f"§\t{base_text} {dots}  {current_page:02d}"
        toc_lines.append(line)

        # Policz strony produktu (zakładam 1-2 strony na produkt, można to poprawić)
        # TODO: Możesz dodać logikę zliczania stron z conversion_cache
        product_path = os.path.join(PRODUKTY_DIR, f'{product_id}.docx')
        file_hash = get_file_hash(product_path)
        if file_hash and file_hash in conversion_cache:
            product_pages = len(conversion_cache[file_hash])
        else:
            product_pages = 1  # Domyślnie 1 strona

        current_page += product_pages

    return '\n'.join(toc_lines)


def inject_toc_into_doc(doc, toc_text):
    """
    Wstawia spis treści do dokumentu DOCX
    Szuka {{SPIS_TRESCI}} lub {{TOC}} i zastępuje tekstem
    """
    injected = False

    for para in doc.paragraphs:
        if '{{SPIS_TRESCI}}' in para.text or '{{TOC}}' in para.text:
            para.text = para.text.replace('{{SPIS_TRESCI}}', toc_text)
            para.text = para.text.replace('{{TOC}}', toc_text)
            injected = True
            print(f"[TOC] Wstawiono spis treści (zastąpiono placeholder)")
            break

    # Jeśli nie ma placeholdera, dodaj na końcu
    if not injected:
        print(f"[TOC] Brak placeholdera - dodaję na końcu dokumentu")
        doc.add_paragraph(toc_text)

    return doc


def get_available_products():
    """Pobiera listę dostępnych produktów z folderu produkty + wykrywa {{placeholders}}"""
    products = []
    if os.path.exists(PRODUKTY_DIR):
        for filename in sorted(os.listdir(PRODUKTY_DIR)):
            if filename.endswith('.docx') and not filename.startswith('~$'):
                product_id = filename.replace('.docx', '')
                product_path = os.path.join(PRODUKTY_DIR, filename)

                # Wykryj placeholders w produkcie
                placeholders = extract_placeholders_from_docx(product_path)

                products.append({
                    'id': product_id,
                    'name': f'Produkt {product_id}',
                    'filename': filename,
                    'placeholders': list(placeholders),  # Konwertuj set na list dla JSON
                    'has_custom_fields': len(placeholders) > 0
                })
    return products


def replace_placeholders(doc, data):
    """Zamienia placeholdery w dokumencie na rzeczywiste dane"""
    # Zamiana w paragrafach
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if key != 'produkty':  # Produkty obsługujemy osobno
                placeholder = '{{' + key + '}}'
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(value))

    # Zamiana w tabelach
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    if key != 'produkty':
                        placeholder = '{{' + key + '}}'
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(value))

    return doc


def find_injection_point(doc, search_text):
    """Znajduje punkt wstrzyknięcia produktów w dokumencie"""
    for i, paragraph in enumerate(doc.paragraphs):
        # Usuń placeholder z tekstu szukanego
        clean_search = search_text.replace('{{opis}}', '')
        if clean_search.strip() in paragraph.text:
            return i
    return None


def insert_page_break(paragraph):
    """Dodaje page break przed paragrafem"""
    run = paragraph.insert_paragraph_before().add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)


def split_document_at_injection_point(template_path, injection_index, data):
    """Dzieli szablon na dwie części: przed i po punkcie wstrzyknięcia

    UWAGA: W większości przypadków strona 1 to WSZYSTKO przed page break,
    więc part2 zazwyczaj będzie None (brak strony 2 w szablonie)
    """

    # Wczytaj szablon z zamienionymi placeholderami
    doc = Document(template_path)
    doc = replace_placeholders(doc, data)

    # Sprawdź czy w szablonie są wyraźne page breaks (section breaks)
    has_explicit_page_break = False
    page_break_at_section = -1

    for i, section in enumerate(doc.sections):
        if i > 0:  # Jest więcej niż jedna sekcja - znaczy że jest page break
            has_explicit_page_break = True
            page_break_at_section = i
            break

    if has_explicit_page_break:
        # Szablon ma wyraźny podział na sekcje (strony)
        # Part1 = pierwsza sekcja, Part2 = reszta
        part1_doc = Document(template_path)
        part1_doc = replace_placeholders(part1_doc, data)

        # TODO: Zaimplementuj podział według sekcji
        # Na razie zwróć cały dokument jako part1
        part2_doc = None
    else:
        # Szablon NIE ma wyraźnego podziału
        # WSZYSTKO jest stroną 1, part2 nie istnieje
        part1_doc = Document(template_path)
        part1_doc = replace_placeholders(part1_doc, data)
        part2_doc = None

    return part1_doc, part2_doc


def add_page_break_to_doc(doc):
    """Dodaje page break na końcu dokumentu"""
    # Dodaj prosty page break na końcu
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)
    return doc


def generate_multifile_offer_docx(data, selected_products, template_data):
    """Generuje ofertę DOCX dla multi-file template (WolfTax)"""
    print(f"[MULTIFILE DOCX] Generuję ofertę z {len(template_data['files'])} plików")

    form_data = data.get('formData', data)  # Kompatybilność
    product_custom_fields = data.get('productCustomFields', {})

    template_folder = os.path.join(TEMPLATES_DIR, template_data['folder'])
    files = sorted(template_data['files'], key=lambda x: x['order'])
    injection_point = template_data.get('injection_point', {})

    # Zacznij od pustego dokumentu
    first_file_path = os.path.join(template_folder, files[0]['file'])
    first_doc = Document(first_file_path)
    first_doc = replace_placeholders(first_doc, form_data)

    composer = Composer(first_doc)

    # Przetwórz pozostałe pliki
    for file_info in files[1:]:
        file_path = os.path.join(template_folder, file_info['file'])

        if not os.path.exists(file_path):
            continue

        doc = Document(file_path)
        doc = replace_placeholders(doc, form_data)

        # SPIS TREŚCI
        if file_info.get('is_toc') and len(selected_products) > 0:
            toc_config = template_data.get('toc', {})
            start_page = toc_config.get('start_page', 5)
            toc_text = generate_table_of_contents(selected_products, product_custom_fields, start_page)
            doc = inject_toc_into_doc(doc, toc_text)
            print(f"[TOC DOCX] Spis treści dodany")

        # INJECTION POINT - produkty
        if (injection_point.get('type') == 'between_files' and
            file_info['file'] == injection_point.get('after')):

            # Najpierw dodaj ten dokument
            composer.append(doc, restart_numbering=True)

            # Wstaw produkty
            print(f"[MULTIFILE DOCX] Wstawiam {len(selected_products)} produktów")
            for product_id in selected_products:
                product_path = os.path.join(PRODUKTY_DIR, f'{product_id}.docx')
                if os.path.exists(product_path):
                    product_doc = Document(product_path)

                    # Custom fields
                    if product_id in product_custom_fields:
                        custom_data = product_custom_fields[product_id]
                        product_doc = replace_placeholders(product_doc, custom_data)

                    product_doc = add_page_break_to_doc(product_doc)
                    composer.append(product_doc, restart_numbering=True)

            continue  # Już dodany

        # Normalnie dodaj dokument
        composer.append(doc, restart_numbering=True)

    # Zapisz
    client_name = form_data.get('KLIENT(NIP)') or form_data.get('klient') or 'Klient'
    output_filename = f"Oferta_WolfTax_{client_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    output_path = os.path.join(GENERATED_OFFERS_DIR, output_filename)

    composer.save(output_path)
    print(f"[MULTIFILE DOCX] Zapisano: {output_filename}")

    return output_path, output_filename


def generate_offer_docx(data, selected_products, template_data=None):
    """Generuje dokument DOCX oferty - MULTI-TEMPLATE SUPPORT"""

    # Sprawdź czy to multi-file czy single-file
    if template_data and template_data.get('type') == 'multi_file':
        return generate_multifile_offer_docx(data, selected_products, template_data)

    # STARY SYSTEM - AIDROPS
    template_config = load_template_config('oferta1')
    template_path = os.path.join(TEMPLATES_DIR, template_config['template_file'])

    # Znajdź punkt wstrzyknięcia
    injection_config = template_config.get('injection_point', {})
    injection_text = injection_config.get('after_paragraph_containing', '')

    # Wczytaj szablon tymczasowo aby znaleźć injection point
    temp_doc = Document(template_path)
    temp_doc = replace_placeholders(temp_doc, data)
    injection_index = find_injection_point(temp_doc, injection_text)

    if injection_index is None:
        # Jeśli nie znaleziono punktu wstrzyknięcia, użyj prostego podejścia
        main_doc = Document(template_path)
        main_doc = replace_placeholders(main_doc, data)

        if selected_products:
            # Dodaj page break na końcu pierwszej strony
            main_doc = add_page_break_to_doc(main_doc)

            composer = Composer(main_doc)
            for product_id in selected_products:
                product_path = os.path.join(PRODUKTY_DIR, f'{product_id}.docx')
                if os.path.exists(product_path):
                    product_doc = Document(product_path)
                    # Dodaj page break na końcu każdego produktu
                    product_doc = add_page_break_to_doc(product_doc)
                    # WAŻNE: restart_numbering=True aby produkt był niezależny
                    composer.append(product_doc, restart_numbering=True)
        else:
            composer = Composer(main_doc)
    else:
        # Podziel szablon na części
        part1_doc, part2_doc = split_document_at_injection_point(template_path, injection_index, data)

        # Dodaj page break na końcu part1
        part1_doc = add_page_break_to_doc(part1_doc)

        # Rozpocznij od pierwszej części
        composer = Composer(part1_doc)

        # Dodaj wszystkie produkty JAKO NIEZALEŻNE DOKUMENTY
        for product_id in selected_products:
            product_path = os.path.join(PRODUKTY_DIR, f'{product_id}.docx')
            if os.path.exists(product_path):
                product_doc = Document(product_path)

                # CUSTOM FIELDS: Jeśli produkt ma custom fields, podstaw wartości
                if product_id in data.get('productCustomFields', {}):
                    custom_data = data['productCustomFields'][product_id]
                    print(f"[DEBUG] Podstawiam custom fields dla produktu {product_id}: {custom_data}")
                    product_doc = replace_placeholders(product_doc, custom_data)

                # Dodaj page break na końcu każdego produktu
                product_doc = add_page_break_to_doc(product_doc)
                # WAŻNE: restart_numbering=True aby każdy produkt zachował swoją sekcję
                composer.append(product_doc, restart_numbering=True)

        # Dodaj drugą część (jeśli istnieje)
        if part2_doc:
            composer.append(part2_doc)

    # Zapisz dokument
    output_filename = f"Oferta_{data.get('KLIENT(NIP)', 'Klient')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    output_path = os.path.join(GENERATED_OFFERS_DIR, output_filename)

    composer.save(output_path)

    return output_path, output_filename


@app.route('/')
def index():
    """Główna strona aplikacji"""
    response = make_response(render_template('index.html'))
    response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    return response


@app.route('/api/templates')
def get_templates():
    """Zwraca listę dostępnych szablonów"""
    templates_config_path = os.path.join(TEMPLATES_DIR, 'templates.json')

    with open(templates_config_path, 'r', encoding='utf-8') as f:
        templates_data = json.load(f)

    return jsonify(templates_data)


@app.route('/api/template/<template_id>')
def get_template_details(template_id):
    """Zwraca szczegóły szablonu + wykrywa wszystkie {{placeholders}}"""
    templates_config_path = os.path.join(TEMPLATES_DIR, 'templates.json')

    with open(templates_config_path, 'r', encoding='utf-8') as f:
        templates_data = json.load(f)

    # Znajdź szablon
    template = None
    for t in templates_data['templates']:
        if t['id'] == template_id:
            template = t
            break

    if not template:
        return jsonify({'error': 'Szablon nie znaleziony'}), 404

    # Skanuj wszystkie pliki i zbierz placeholders
    all_placeholders = {}

    if template['type'] == 'single_file':
        # Stary system - pojedynczy plik
        file_path = os.path.join(TEMPLATES_DIR, template.get('folder', '.'), template['main_file'])
        placeholders = extract_placeholders_from_docx(file_path)
        all_placeholders['main'] = list(placeholders)

    elif template['type'] == 'multi_file':
        # Nowy system - wiele plików
        template_folder = os.path.join(TEMPLATES_DIR, template['folder'])

        for file_info in template['files']:
            file_path = os.path.join(template_folder, file_info['file'])
            if os.path.exists(file_path):
                placeholders = extract_placeholders_from_docx(file_path)
                if placeholders:
                    all_placeholders[file_info['file']] = {
                        'name': file_info.get('name', file_info['file']),
                        'placeholders': list(placeholders)
                    }

    template['discovered_placeholders'] = all_placeholders
    template['total_placeholders'] = sum(
        len(v['placeholders']) if isinstance(v, dict) else len(v)
        for v in all_placeholders.values()
    )

    return jsonify(template)


@app.route('/api/template-config')
def get_template_config():
    """Zwraca konfigurację szablonu (legacy compatibility)"""
    config = load_template_config('oferta1')
    return jsonify(config)


@app.route('/api/products')
def get_products():
    """Zwraca listę dostępnych produktów"""
    products = get_available_products()
    return jsonify(products)


@app.route('/api/save-offer', methods=['POST'])
def save_offer():
    """Zapisuje ofertę do pliku JSON"""
    data = request.json
    offer_name = data.get('offer_name', f"oferta_{datetime.now().strftime('%Y%m%d_%H%M%S')}")

    # Usuń offer_name z danych do zapisu
    offer_data = {k: v for k, v in data.items() if k != 'offer_name'}

    filename = f"{offer_name}.json"
    filepath = os.path.join(SAVED_OFFERS_DIR, filename)

    with open(filepath, 'w', encoding='utf-8') as f:
        json.dump(offer_data, f, ensure_ascii=False, indent=2)

    return jsonify({'success': True, 'filename': filename})


@app.route('/api/load-offer/<filename>')
def load_offer(filename):
    """Wczytuje zapisaną ofertę"""
    filepath = os.path.join(SAVED_OFFERS_DIR, filename)

    if not os.path.exists(filepath):
        return jsonify({'error': 'Plik nie istnieje'}), 404

    with open(filepath, 'r', encoding='utf-8') as f:
        data = json.load(f)

    return jsonify(data)


@app.route('/api/saved-offers')
def get_saved_offers():
    """Zwraca listę zapisanych ofert"""
    offers = []
    if os.path.exists(SAVED_OFFERS_DIR):
        for filename in os.listdir(SAVED_OFFERS_DIR):
            if filename.endswith('.json'):
                filepath = os.path.join(SAVED_OFFERS_DIR, filename)
                stat = os.stat(filepath)
                offers.append({
                    'filename': filename,
                    'name': filename.replace('.json', ''),
                    'modified': datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M:%S')
                })

    # Sortuj po dacie modyfikacji (najnowsze pierwsze)
    offers.sort(key=lambda x: x['modified'], reverse=True)

    return jsonify(offers)


@app.route('/api/generate-offer', methods=['POST'])
def generate_offer():
    """Generuje dokument DOCX oferty z progress tracking"""
    import time
    start_time = time.time()

    data = request.json
    form_data = data.get('formData', {})
    selected_products = data.get('selectedProducts', [])
    template_data = data.get('templateData')  # Obsługa multi-template

    try:
        # Progress: Start
        send_progress("⚙️ Rozpoczynam generowanie dokumentu DOCX...", 10)

        # Progress: Tworzenie szablonu
        if template_data and template_data.get('type') == 'multi_file':
            send_progress(f"📝 Przetwarzam szablon {template_data.get('name', 'WolfTax')}...", 20)
        else:
            send_progress("📝 Przetwarzam szablon oferty...", 20)

        # Progress: Dodawanie produktów
        total_products = len(selected_products)
        if total_products > 0:
            send_progress(f"📦 Dodaję {total_products} produktów...", 40)

        output_path, output_filename = generate_offer_docx(data, selected_products, template_data)

        # Progress: Zapisywanie
        send_progress("💾 Zapisuję dokument...", 80)

        # Progress: Gotowe
        elapsed_time = time.time() - start_time
        send_progress(f"✅ Dokument gotowy! ({elapsed_time:.1f}s)", 100)

        # Krótka pauza aby pokazać 100%
        time.sleep(0.3)

        # Ukryj pasek
        send_progress("", 0)

        return jsonify({
            'success': True,
            'filename': output_filename,
            'download_url': f'/api/download-offer/{output_filename}',
            'generation_time': f"{elapsed_time:.2f}s"
        })
    except Exception as e:
        send_progress(f"❌ Błąd: {str(e)}", 0)
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/download-offer/<filename>')
def download_offer(filename):
    """Pobiera wygenerowaną ofertę"""
    filepath = os.path.join(GENERATED_OFFERS_DIR, filename)

    if not os.path.exists(filepath):
        return jsonify({'error': 'Plik nie istnieje'}), 404

    return send_file(filepath, as_attachment=True, download_name=filename)


@app.route('/api/preview-full-offer', methods=['POST'])
def preview_full_offer():
    """
    Generuje podgląd - INTELIGENTNY CACHE SYSTEM!

    Strategia:
    1. Sprawdź INTELIGENTNY CACHE (kompletna oferta)
    2. Jeśli HIT → wyślij NATYCHMIAST (<200ms!)
    3. Jeśli MISS → generuj i zapisz do cache
    """
    data = request.json
    template_id = data.get('templateId', 'aidrops')
    template_data = data.get('templateData')
    form_data = data.get('formData', {})
    selected_products = data.get('selectedProducts', [])
    product_custom_fields = data.get('productCustomFields', {})
    use_streaming = data.get('streaming', True)

    print(f"[PREVIEW] Template: {template_id}, Produkty: {selected_products}")

    # ============================================================
    # INTELIGENTNY CACHE - sprawdź czy ta DOKŁADNA oferta już istnieje
    # ============================================================
    cache_key = get_offer_cache_key(template_id, form_data, selected_products, product_custom_fields)

    if cache_key and cache_key in offer_cache:
        print(f"[CACHE] 🎯 INTELIGENTNY CACHE HIT! Key: {cache_key[:16]}...")
        send_progress("⚡ Ładuję z cache...", 10)

        cached_offer = offer_cache[cache_key]
        pages_metadata = cached_offer['pages']
        total_pages = cached_offer['total_pages']

        print(f"[CACHE] ✓ Wysyłam {total_pages} stron z cache (INSTANT!)")

        # Wyślij wszystkie strony NATYCHMIAST przez WebSocket
        if use_streaming:
            for page in pages_metadata:
                send_page_ready(page)

        # Usuń obrazy z metadanych dla HTTP response (zmniejsz rozmiar)
        metadata_without_images = []
        for meta in pages_metadata:
            meta_copy = {k: v for k, v in meta.items() if k != 'image'}
            meta_copy['has_image'] = True
            metadata_without_images.append(meta_copy)

        send_progress("✅ Gotowe!", 100)

        return jsonify({
            'success': True,
            'total_pages': total_pages,
            'pages_metadata': metadata_without_images,
            'from_cache': True
        })

    # CACHE MISS - generuj normalnie
    print(f"[CACHE] ❌ MISS - generuję nową ofertę... Key: {cache_key[:16] if cache_key else 'None'}...")
    send_progress("Generuję podgląd...", 5)

    pages_metadata = []
    page_counter = 0

    # Wybierz strategię na podstawie typu szablonu
    if template_data and template_data.get('type') == 'multi_file':
        # NOWY SYSTEM - WolfTax (wiele plików)
        print("[DEBUG] Używam multi-file template")
        send_progress("Ładuję szablon z cache...", 10)

        template_folder = os.path.join(TEMPLATES_DIR, template_data['folder'])
        files = sorted(template_data['files'], key=lambda x: x['order'])
        injection_point = template_data.get('injection_point', {})

        # Procesuj każdy plik (ON-DEMAND, bez pre-cache!)
        for file_info in files:
            file_name = file_info['file']
            file_path = os.path.join(template_folder, file_name)

            if not os.path.exists(file_path):
                print(f"[WARNING] Plik nie istnieje: {file_path}")
                continue

            print(f"[PREVIEW] Przetwarzam: {file_name}")
            send_progress(f"📄 {file_info.get('name', file_name)}...", 10 + page_counter * 2)

            # Generuj plik (zawsze z aktualnymi danymi)
            # Załaduj i wypełnij placeholders
            doc = Document(file_path)
            doc = replace_placeholders(doc, form_data)

            # SPIS TREŚCI - jeśli to plik TOC, wygeneruj i wstaw
            if needs_toc:
                toc_config = template_data.get('toc', {})
                start_page = toc_config.get('start_page', 5)

                print(f"[TOC] Generuję spis treści dla {len(selected_products)} produktów (start: strona {start_page})")
                toc_text = generate_table_of_contents(selected_products, product_custom_fields, start_page)
                doc = inject_toc_into_doc(doc, toc_text)
                print(f"[TOC] Spis treści wygenerowany:\n{toc_text}")

            # Konwertuj na obrazy
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx', dir=PREVIEW_CACHE_DIR)
            doc.save(temp_file.name)
            temp_file.close()

            file_images = convert_docx_to_images(temp_file.name, use_cache=False, progress_callback=None)

            try:
                os.unlink(temp_file.name)
            except:
                pass

            # Wyślij strony tego pliku
            for idx, img_data in enumerate(file_images):
                page_counter += 1
                page_data = {
                    'type': 'template',
                    'number': page_counter,
                    'image': img_data,
                    'has_image': True,
                    'page_index': idx,
                    'status': 'ready',
                    'source_file': file_name
                }
                pages_metadata.append(page_data)

                if use_streaming:
                    send_page_ready(page_data)

            # INJECTION POINT - wstaw produkty po tym pliku
            if (injection_point.get('type') == 'between_files' and
                file_info['file'] == injection_point.get('after')):

                print(f"[DEBUG] Injection point! Wstawiam {len(selected_products)} produktów po {file_info['file']}")
                send_progress(f"Dodaję produkty...", 50)

                # Wstaw produkty w tym miejscu (synchronicznie dla prostoty)
                for product_idx, product_id in enumerate(selected_products):
                    product_path = os.path.join(PRODUKTY_DIR, f'{product_id}.docx')

                    if os.path.exists(product_path):
                        print(f"[DEBUG] Dodaję produkt: {product_id}")

                        # Custom fields
                        path_to_convert = product_path
                        temp_file = None

                        if product_id in product_custom_fields and product_custom_fields[product_id]:
                            custom_data = product_custom_fields[product_id]
                            print(f"[CUSTOM FIELDS] Podstawiam dla {product_id}: {custom_data}")

                            product_doc = Document(product_path)
                            product_doc = replace_placeholders(product_doc, custom_data)

                            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx', dir=PREVIEW_CACHE_DIR)
                            product_doc.save(temp_file.name)
                            temp_file.close()
                            path_to_convert = temp_file.name

                        # Konwertuj produkt
                        use_cache = (temp_file is None)
                        product_images = convert_docx_to_images(path_to_convert, use_cache=use_cache, progress_callback=None)

                        # Cleanup
                        if temp_file:
                            try:
                                os.unlink(temp_file.name)
                            except:
                                pass

                        # Wyślij strony produktu
                        for idx, img_data in enumerate(product_images):
                            page_counter += 1
                            page_data = {
                                'type': 'product',
                                'number': page_counter,
                                'product_id': product_id,
                                'image': img_data,
                                'has_image': True,
                                'page_index': idx,
                                'status': 'ready'
                            }
                            pages_metadata.append(page_data)

                            if use_streaming:
                                send_page_ready(page_data)

                        print(f"[DEBUG] ✓ Produkt {product_id} dodany ({len(product_images)} stron)")

    else:
        # AIDROPS (pojedynczy plik) - używa INTELIGENTNY CACHE
        print("[DEBUG] Używam single-file template (AIDROPS)")
        send_progress("Generuję szablon...", 10)

        # Generuj szablon z aktualnymi danymi
        template_config = load_template_config('oferta1')
        template_path = os.path.join(TEMPLATES_DIR, template_config['template_file'])

        doc = Document(template_path)
        doc = replace_placeholders(doc, form_data)

        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx', dir=PREVIEW_CACHE_DIR)
        doc.save(temp_file.name)
        temp_file.close()

        template_images = convert_docx_to_images(temp_file.name, use_cache=False, progress_callback=send_progress)

        try:
            os.unlink(temp_file.name)
        except:
            pass

        # Wyślij strony szablonu
        for idx, img_data in enumerate(template_images):
            page_counter += 1
            page_data = {
                'type': 'template',
                'number': page_counter,
                'image': img_data,
                'has_image': True,
                'page_index': idx,
                'status': 'ready'
            }
            pages_metadata.append(page_data)

            if use_streaming:
                send_page_ready(page_data)

    send_progress(f"Szablon gotowy ({page_counter} stron)", 40)

    # === PRODUKTY (tylko dla single-file / AIDROPS) ===
    # Dla multi-file produkty są już dodane w injection point powyżej!

    if template_data and template_data.get('type') == 'multi_file':
        # Multi-file - produkty już dodane w injection point
        print("[DEBUG] Multi-file: produkty już wstawione w injection point")
        send_progress("Wszystkie strony gotowe!", 100)

        import time
        time.sleep(0.3)
        send_progress("", 0)

        print(f"[DEBUG] Streaming zakończony - {len(pages_metadata)} stron")

        # WAŻNE: Usuń obrazy z metadanych (są już wysłane przez WebSocket!)
        metadata_without_images = []
        for meta in pages_metadata:
            meta_copy = {k: v for k, v in meta.items() if k != 'image'}
            meta_copy['has_image'] = meta.get('status') == 'ready'
            metadata_without_images.append(meta_copy)

        return jsonify({
            'success': True,
            'total_pages': len(pages_metadata),
            'pages_metadata': metadata_without_images
        })

    # === AIDROPS - STARE PRODUKTY Z PARALLEL PROCESSING ===
    total_products = len(selected_products)

    # Najpierw wyślij metadane wszystkich produktów (wyszarzone)
    for product_idx, product_id in enumerate(selected_products):
        product_path = os.path.join(PRODUKTY_DIR, f'{product_id}.docx')

        if os.path.exists(product_path):
            # Sprawdź ile stron ma produkt (z cache lub szybkie sprawdzenie)
            file_hash = get_file_hash(product_path)
            if file_hash and file_hash in conversion_cache:
                product_images_count = len(conversion_cache[file_hash])
            else:
                product_images_count = 1  # Domyślnie 1 strona

            for page_idx in range(product_images_count):
                page_counter += 1
                pages_metadata.append({
                    'type': 'product',
                    'number': page_counter,
                    'product_id': product_id,
                    'page_index': page_idx,
                    'status': 'pending'  # Wyszarzona - oczekuje na generowanie
                })

    # PARALLEL PROCESSING: Generuj produkty równolegle!
    current_page = page_counter + 1

    def process_product(product_idx, product_id):
        """Przetwórz pojedynczy produkt (thread-safe)"""
        product_path = os.path.join(PRODUKTY_DIR, f'{product_id}.docx')

        if not os.path.exists(product_path):
            return None

        # Wyślij status "generating"
        if use_streaming:
            socketio.emit('product_status', {
                'product_id': product_id,
                'status': 'generating'
            })

        print(f"[PARALLEL] Start konwersji produktu {product_id}")

        # CUSTOM FIELDS: Jeśli produkt ma custom fields, stwórz tymczasową kopię z wypełnionymi wartościami
        path_to_convert = product_path
        temp_file = None

        if product_id in product_custom_fields and product_custom_fields[product_id]:
            custom_data = product_custom_fields[product_id]
            print(f"[CUSTOM FIELDS] Podstawiam dla produktu {product_id}: {custom_data}")

            # Załaduj produkt i wypełnij placeholders
            product_doc = Document(product_path)
            product_doc = replace_placeholders(product_doc, custom_data)

            # Zapisz do tymczasowego pliku
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx', dir=PREVIEW_CACHE_DIR)
            product_doc.save(temp_file.name)
            temp_file.close()
            path_to_convert = temp_file.name

        # Konwertuj (cache działa wewnątrz funkcji - ale custom fields maja cache wyłączony)
        use_cache_for_product = (temp_file is None)  # Cache tylko gdy NIE ma custom fields
        product_images = convert_docx_to_images(path_to_convert, use_cache=use_cache_for_product, progress_callback=None)

        # Usuń tymczasowy plik
        if temp_file:
            try:
                os.unlink(temp_file.name)
            except:
                pass

        print(f"[PARALLEL] Zakończono konwersję produktu {product_id}: {len(product_images)} stron")

        return {
            'product_id': product_id,
            'product_idx': product_idx,
            'images': product_images
        }

    # Użyj ThreadPoolExecutor do przetwarzania równoległego (max 3 produkty naraz)
    if len(selected_products) > 0:
        max_workers = min(3, len(selected_products))

        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            # Uruchom wszystkie produkty równolegle
            futures = {
                executor.submit(process_product, idx, pid): pid
                for idx, pid in enumerate(selected_products)
            }

            completed_count = 0

            # Odbieraj wyniki w miarę gotowości
            for future in as_completed(futures):
                product_id = futures[future]

                try:
                    result = future.result()

                    if result:
                        completed_count += 1
                        progress_percent = 30 + int(65 * (completed_count / max(total_products, 1)))
                        send_progress(f"Produkt {completed_count}/{total_products} gotowy...", progress_percent)

                        # STREAMING: Wysyłaj strony przez WebSocket
                        for idx, img_data in enumerate(result['images']):
                            # Oblicz numer strony (produkty mogą kończyć się w różnej kolejności!)
                            page_num = current_page + result['product_idx']

                            page_data = {
                                'type': 'product',
                                'number': page_num,
                                'product_id': result['product_id'],
                                'image': img_data,
                                'has_image': True,
                                'page_index': idx,
                                'status': 'ready'
                            }

                            if use_streaming:
                                send_page_ready(page_data)

                        print(f"[PARALLEL] ✓ Produkt {product_id} wysłany przez WebSocket")

                except Exception as e:
                    print(f"[ERROR] Błąd przetwarzania produktu {product_id}: {e}")
                    import traceback
                    traceback.print_exc()

    send_progress("Wszystkie strony gotowe!", 100)

    # Daj chwilę na wyświetlenie 100%
    import time
    time.sleep(0.3)

    send_progress("", 0)  # Ukryj pasek

    print(f"[DEBUG] Streaming zakończony - {len(pages_metadata)} stron")

    # WAŻNE: Usuń obrazy z metadanych (są już wysłane przez WebSocket!)
    metadata_without_images = []
    for meta in pages_metadata:
        meta_copy = {k: v for k, v in meta.items() if k != 'image'}
        meta_copy['has_image'] = meta.get('status') == 'ready'
        metadata_without_images.append(meta_copy)

    # ============================================================
    # ZAPISZ DO INTELIGENTNEGO CACHE!
    # ============================================================
    if cache_key:
        import time
        offer_cache[cache_key] = {
            'pages': pages_metadata,  # PEŁNE dane z obrazami!
            'total_pages': len(pages_metadata),
            'timestamp': time.time()
        }
        print(f"[CACHE] ✓ Zapisano do cache: {cache_key[:16]}... ({len(pages_metadata)} stron)")
        print(f"[CACHE] Cache size: {len(offer_cache)} ofert")

    return jsonify({
        'success': True,
        'streaming': use_streaming,
        'pages_metadata': metadata_without_images,
        'total_pages': len(pages_metadata),
        'from_cache': False
    })


@app.route('/api/load-page', methods=['POST'])
def load_single_page():
    """Lazy loading: ładuje pojedynczą stronę na żądanie"""
    data = request.json
    page_type = data.get('type')
    page_index = data.get('page_index')
    product_id = data.get('product_id')
    form_data = data.get('formData', {})

    print(f"[DEBUG] Lazy loading strony: type={page_type}, index={page_index}, product={product_id}")

    try:
        if page_type == 'template':
            # Sprawdź cache szablonu
            form_hash = get_form_data_hash(form_data)
            if form_hash and form_hash in template_cache:
                template_images = template_cache[form_hash]
                if 0 <= page_index < len(template_images):
                    return jsonify({
                        'success': True,
                        'image': template_images[page_index]
                    })

        elif page_type == 'product' and product_id:
            # Pobierz z cache produktu
            product_path = os.path.join(PRODUKTY_DIR, f'{product_id}.docx')
            if os.path.exists(product_path):
                product_images = convert_docx_to_images(product_path, use_cache=True)
                if 0 <= page_index < len(product_images):
                    return jsonify({
                        'success': True,
                        'image': product_images[page_index]
                    })

        return jsonify({'success': False, 'error': 'Strona nie znaleziona'})

    except Exception as e:
        print(f"[ERROR] Błąd lazy loading: {e}")
        return jsonify({'success': False, 'error': str(e)})


# WebSocket event handlers
@socketio.on('connect')
def handle_connect():
    print('=' * 80)
    print(f'[WebSocket] ✅ Klient połączony! SID: {request.sid}')
    print('=' * 80)

@socketio.on('disconnect')
def handle_disconnect():
    print(f'[WebSocket] ❌ Klient rozłączony: {request.sid}')

# ============================================================
# STARTUP: Uruchom unoserver i pre-loading
# ============================================================
print("=" * 80)
print("[STARTUP] 🚀 Inicjalizacja systemu...")
print("=" * 80)

# KROK 1: Uruchom unoserver jeśli nie działa (tylko na serwerze)
start_unoserver()

# KROK 2: Pre-loading produktów i szablonów w tle
print("[STARTUP] Uruchamiam pre-loading w tle...")
preload_products_async()
preload_templates_async()

print("[STARTUP] ✅ System gotowy do pracy!")
print("=" * 80)

if __name__ == '__main__':
    socketio.run(app, debug=True, host='0.0.0.0', port=40207, allow_unsafe_werkzeug=True)
